import os
import logging
from typing import List, Dict, Optional
from fastapi import UploadFile
from sqlalchemy.orm import Session
import PyPDF2
from docx import Document as DocxDocument
from io import BytesIO

from ..models.document import Document
from .metadata_extractor import MetadataExtractor

logger = logging.getLogger(__name__)

class DocumentService:
    """Service for handling document uploads and processing"""
    
    def __init__(self):
        self.metadata_extractor = MetadataExtractor()
    
    async def process_upload(self, files: List[UploadFile], db: Session) -> Dict[str, int]:
        """Process multiple uploaded files"""
        processed = 0
        failed = 0
        
        for file in files:
            try:
                await self._process_single_file(file, db)
                processed += 1
                logger.info(f"Successfully processed {file.filename}")
            except Exception as e:
                failed += 1
                logger.error(f"Failed to process {file.filename}: {e}")
        
        return {"processed": processed, "failed": failed}
    
    async def _process_single_file(self, file: UploadFile, db: Session):
        """Process a single uploaded file"""
        # Validate file type
        if not self._is_valid_file_type(file.filename):
            raise ValueError(f"Unsupported file type: {file.filename}")
        
        # Read file content
        content = await file.read()
        
        # Extract text based on file type
        text_content = self._extract_text(content, file.filename)
        
        # Extract metadata
        metadata = self.metadata_extractor.extract_metadata(text_content, file.filename)
        
        # Create document record
        document = Document(
            filename=file.filename,
            file_type=self._get_file_type(file.filename),
            file_size=len(content),
            content=text_content,
            **metadata
        )
        
        # Save to database
        db.add(document)
        db.commit()
        db.refresh(document)
    
    def _is_valid_file_type(self, filename: str) -> bool:
        """Check if file type is supported"""
        if not filename:
            return False
        
        extension = filename.lower().split('.')[-1]
        return extension in ['pdf', 'docx']
    
    def _get_file_type(self, filename: str) -> str:
        """Get file type from filename"""
        return filename.lower().split('.')[-1]
    
    def _extract_text(self, content: bytes, filename: str) -> str:
        """Extract text from PDF or DOCX file"""
        file_type = self._get_file_type(filename)
        
        if file_type == 'pdf':
            return self._extract_pdf_text(content)
        elif file_type == 'docx':
            return self._extract_docx_text(content)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_file = BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            raise
    
    def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            docx_file = BytesIO(content)
            doc = DocxDocument(docx_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            raise
    
    def get_all_documents(self, db: Session) -> List[Document]:
        """Get all documents from database"""
        return db.query(Document).all()
    
    def get_document_by_id(self, document_id: int, db: Session) -> Optional[Document]:
        """Get document by ID"""
        return db.query(Document).filter(Document.id == document_id).first()
