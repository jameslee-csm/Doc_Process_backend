#!/usr/bin/env python3
"""
Demo data generator for Legal Intel Dashboard
Creates sample documents with metadata for testing
"""

import os
import sys
from io import BytesIO
from docx import Document as DocxDocument
from PyPDF2 import PdfWriter, PdfReader
import tempfile

# Add the app directory to the Python path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'app'))

from app.models.database import SessionLocal, init_db
from app.models.document import Document
from app.services.metadata_extractor import MetadataExtractor

def create_sample_docx(filename, content):
    """Create a sample DOCX file with given content"""
    doc = DocxDocument()
    doc.add_paragraph(content)
    
    # Save to temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    
    # Read the file content
    with open(temp_file.name, 'rb') as f:
        content_bytes = f.read()
    
    # Clean up
    os.unlink(temp_file.name)
    
    return content_bytes

def create_sample_pdf(filename, content):
    """Create a sample PDF file with given content"""
    # For simplicity, we'll create a minimal PDF
    # In a real scenario, you'd use a proper PDF library
    writer = PdfWriter()
    
    # Create a temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    
    # Write empty PDF (this is a simplified version)
    writer.write(temp_file)
    temp_file.close()
    
    # Read the file content
    with open(temp_file.name, 'rb') as f:
        content_bytes = f.read()
    
    # Clean up
    os.unlink(temp_file.name)
    
    return content_bytes

def generate_demo_data():
    """Generate demo documents with metadata"""
    print("Generating demo data...")
    
    # Initialize database
    init_db()
    
    # Sample documents with metadata
    demo_documents = [
        {
            'filename': 'nda_technology_uae.docx',
            'content': 'This Non-Disclosure Agreement is governed by UAE law and covers technology services in the Middle East region.',
            'file_type': 'docx'
        },
        {
            'filename': 'msa_oil_gas_uk.pdf',
            'content': 'Master Service Agreement for oil and gas exploration services governed by UK law in Europe.',
            'file_type': 'pdf'
        },
        {
            'filename': 'franchise_agreement_us.docx',
            'content': 'Franchise Agreement for retail operations in North America governed by US law.',
            'file_type': 'docx'
        },
        {
            'filename': 'employment_contract_singapore.pdf',
            'content': 'Employment Agreement for technology company in Singapore, Asia-Pacific region.',
            'file_type': 'pdf'
        },
        {
            'filename': 'license_agreement_hong_kong.docx',
            'content': 'License Agreement for software licensing in Hong Kong, Asia region.',
            'file_type': 'docx'
        },
        {
            'filename': 'service_agreement_qatar.pdf',
            'content': 'Service Agreement for healthcare services in Qatar, Middle East region.',
            'file_type': 'pdf'
        }
    ]
    
    # Initialize services
    metadata_extractor = MetadataExtractor()
    db = SessionLocal()
    
    try:
        for doc_data in demo_documents:
            print(f"Processing {doc_data['filename']}...")
            
            # Create file content based on type
            if doc_data['file_type'] == 'docx':
                file_content = create_sample_docx(doc_data['filename'], doc_data['content'])
            else:
                file_content = create_sample_pdf(doc_data['filename'], doc_data['content'])
            
            # Extract metadata
            metadata = metadata_extractor.extract_metadata(doc_data['content'], doc_data['filename'])
            
            # Create document record
            document = Document(
                filename=doc_data['filename'],
                file_type=doc_data['file_type'],
                file_size=len(file_content),
                content=doc_data['content'],
                **metadata
            )
            
            # Save to database
            db.add(document)
            db.commit()
            
            print(f"  ✓ Added to database with metadata: {metadata}")
        
        print(f"\n✓ Successfully created {len(demo_documents)} demo documents")
        print("You can now test the query and dashboard functionality!")
        
    except Exception as e:
        print(f"Error generating demo data: {e}")
        db.rollback()
    finally:
        db.close()

if __name__ == "__main__":
    generate_demo_data()
